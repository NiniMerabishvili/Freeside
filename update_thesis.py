import docx

doc = docx.Document('Final Thesis.docx')

replacements = [
    ("Six students, six working professionals", "Six students, four working professionals"),
    ("three entrepreneurs, two employed professionals, and one freelancer", "one entrepreneur, two employed professionals, and one freelancer"),
    ("Seven identified as female, five as male", "Four identified as female, six as male"),
    ("Nine debrief interviews", "Eight debrief interviews"),
    ("Eight of nine interviewees", "Eight of eight interviewees"),
    ("t(11)", "t(9)"),
    ("66.7%", "80%"),
    ("41.7%", "50%"),
    ("Anthropic Claude API", "Gemini API"),
    ("Gemini 3.5 Flash", "Gemini 2.5 Flash"),
    ("Next.js 14", "Next.js 16.2.7"),
    ("Conventional energy sliders ask users to accurately introspect and report their own cognitive state. Humans are demonstrably unreliable at that task (Van Dongen et al., 2003). Rather than persist with that broken model, the engine replaces", "The engine replaces")
]

not_found = []

for old, new in replacements:
    found = False
    for p in doc.paragraphs:
        if old in p.text:
            # Try to find it in a single run
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    found = True
            
            # If not found in a single run, we have to do something more complex or just accept formatting loss for this paragraph
            if not found:
                print(f"Warning: '{old}' spans multiple runs in paragraph. Doing paragraph-level replacement (may lose formatting).")
                # Paragraph-level replacement
                # Save formatting of runs? Too complex. Just replace and clear.
                new_text = p.text.replace(old, new)
                p.clear()
                p.add_run(new_text)
                found = True
    if not found:
        not_found.append(old)

if not_found:
    print("Could not find the following strings to replace:")
    for nf in not_found:
        print(f" - {nf}")

doc.save('Final Thesis - Revised.docx')
print("Saved to 'Final Thesis - Revised.docx'")

